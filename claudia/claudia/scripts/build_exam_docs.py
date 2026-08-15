from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("deliverables/PROMESA_Submission")
NAVY, TEAL, MUTED, LIGHT = "172033", "0F766E", "475569", "F1F5F9"

def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.name = "Calibri"; r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def widths(table, values):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, values): cell.width = Inches(width)

def doc(title, subtitle="CSCD 602 Advanced Software Engineering Project"):
    d = Document(); s = d.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    s.header_distance = s.footer_distance = Inches(.49)
    normal = d.styles["Normal"]; normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [("Heading 1",16,"2E74B5",16,8),("Heading 2",13,"2E74B5",12,6),("Heading 3",12,"1F4D78",8,4)]:
        st=d.styles[name]; st.font.name="Calibri"; st._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    hp=s.header.paragraphs[0]; hp.text="PROMESA | CSCD 602 Individual Project"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; hp.runs[0].font.size=Pt(8); hp.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
    fp=s.footer.paragraphs[0]; fp.text="University of Ghana | Department of Computer Science"; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.runs[0].font.size=Pt(8); fp.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(42); p.paragraph_format.space_after=Pt(8)
    r=p.add_run("PROMESA"); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(24); r=p.add_run(title); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(TEAL)
    meta=d.add_table(rows=4, cols=2); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(meta,[1.55,4.95])
    for row,(a,b) in zip(meta.rows,[("Course",subtitle),("Student Name","[Enter your name]"),("Student ID","[Enter your student ID]"),("Submission Date","[Enter date]")]): set_cell(row.cells[0],a,True,LIGHT); set_cell(row.cells[1],b)
    d.add_paragraph(); return d

def h(d,text,level=1): d.add_heading(text,level)
def p(d,text,italic=False): q=d.add_paragraph(); r=q.add_run(text); r.italic=italic; return q
def bullets(d,items):
    for item in items: d.add_paragraph(item, style="List Bullet")
def numbered(d,items):
    for item in items: d.add_paragraph(item, style="List Number")
def table(d, headers, rows, col_widths=None):
    t=d.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if col_widths: widths(t,col_widths)
    for c,x in zip(t.rows[0].cells,headers): set_cell(c,x,True,LIGHT)
    for row in rows:
        cells=t.add_row().cells
        for c,x in zip(cells,row): set_cell(c,x)
    d.add_paragraph(); return t

def save(d,name): OUT.mkdir(parents=True,exist_ok=True); d.save(OUT/name)

def project_documentation():
    d=doc("Consolidated Project Documentation")
    h(d,"1. Project title and problem statement")
    p(d,"PROMESA is a web-based donation campaign management system for SDG-focused NGOs. Small NGOs often manage campaigns through spreadsheets, social media and email, which reduces visibility of donor commitments and makes it difficult to show how work aligns with the UN Sustainable Development Goals.")
    h(d,"2. Aim, objectives and stakeholders")
    p(d,"The aim is to provide a deployable web application for NGO campaign management and donor pledge tracking.")
    bullets(d,["Allow donors to discover approved campaigns and record GHS pledges.","Allow NGOs to submit profiles and campaign proposals for review.","Allow administrators to approve or reject NGOs and campaigns.","Provide role-based dashboards, validation and traceable pledge records."])
    table(d,["Stakeholder","Need"],[["Public visitor","Browse approved campaigns without an account."],["Donor","Pledge support and review pledge history."],["NGO representative","Create an NGO profile, submit campaigns and view received pledges."],["Administrator","Oversee accounts and approve/reject submissions."]],[2.0,4.5])
    h(d,"3. Lifecycle, architecture and design")
    p(d,"The project follows an iterative lifecycle: requirements and estimation, analysis and design, implementation, testing, deployment and planned maintenance. Next.js provides the App Router user interface and server actions. Auth.js provides credentials sessions; Drizzle ORM accesses Neon PostgreSQL; Vercel hosts the application.")
    table(d,["Entity","Key relationships"],[["User","Has one role: donor, NGO or administrator."],["NGO profile","Owned by one NGO user; has an approval state."],["Campaign","Owned by an NGO profile; has a target, deadline, SDG and lifecycle state."],["Pledge","Connects one donor to one campaign and records amount/date/status."]],[1.6,4.9])
    h(d,"4. Effort estimation")
    p(d,"Wideband expert estimation was selected because PROMESA is an individual, time-boxed project with no dependable historic velocity. The estimate was used to retain only core functionality within the 48-hour examination scope.")
    table(d,["Work package","Estimated hours"],[["Requirements and design","8"],["Database, authentication and foundation","9"],["Core workflows","13"],["Testing and refinement","6"],["Deployment","4"],["Documentation","8"],["Total","48"]],[4.8,1.7])
    h(d,"5. Implementation and deployment")
    p(d,"Implemented features include public campaign search and SDG filtering; registration and credentials sign-in; NGO and campaign review; GHS pledge recording; donor, NGO and administrator dashboards; Zod validation; bcrypt password hashing; and database referential integrity.")
    p(d,"Deployment target: Vercel for the Next.js application and Neon for PostgreSQL. Environment variables required are DATABASE_URL, NEXTAUTH_URL, NEXTAUTH_SECRET, ADMIN_EMAIL and ADMIN_PASSWORD.")
    h(d,"6. Limitations, maintenance and future evolution")
    p(d,"The MVP records pledges only and does not collect payments. Campaign coordinator updates, notifications, audit logs, document uploads, multi-currency support and advanced reporting are deferred. Maintenance includes corrective fixes, adaptive dependency updates, perfective accessibility/reporting work, and preventive backups, monitoring and credential rotation.")
    d.add_page_break()
    h(d,"7. References")
    bullets(d,["University of Ghana, CSCD 602 Advanced Software Engineering Project Examination, 2025/2026.","United Nations, Sustainable Development Goals, https://sdgs.un.org/goals.","Next.js Documentation, https://nextjs.org/docs.","Neon Documentation, https://neon.com/docs."])
    save(d,"Project_Documentation.docx")

def srs():
    d=doc("Software Requirements Specification (SRS)")
    h(d,"1. Purpose and scope"); p(d,"PROMESA supports transparent SDG donation campaigns. It is a pledge-management system: it records commitments in Ghana cedi and does not process payment transactions.")
    h(d,"2. Functional requirements")
    rows=[["FR-01","Users can register with name, email, password and donor/NGO role."],["FR-02","Registered users can sign in and sign out using credentials."],["FR-03","Passwords are stored as secure bcrypt hashes; routes/actions enforce roles."],["FR-04","NGO users can submit one profile with contact details and primary SDG."],["FR-05","Administrators can approve or reject pending NGO profiles."],["FR-06","Approved NGOs can submit campaigns with title, description, target, deadline and SDG."],["FR-07","Administrators can approve or reject pending campaigns."],["FR-08","Public users can search approved active campaigns by title and SDG."],["FR-09","Donors can record positive GHS pledges and view pledge history."],["FR-10","NGOs can view received pledges and acknowledge or fulfil them."],["FR-11","Administrators can view recent users and platform submission statistics."]]
    table(d,["ID","Requirement"],rows,[.85,5.65])
    h(d,"3. Non-functional requirements")
    table(d,["Area","Requirement"],[["Usability","Clear responsive layout using solid neutral surfaces and accessible contrast; no decorative gradients."],["Security","Bcrypt passwords, Auth.js sessions, server-side role checks and Zod validation."],["Performance","Campaign pages should load promptly for a small initial dataset."],["Reliability","Vercel-hosted live URL remains available for assessment."],["Maintainability","TypeScript, component separation, typed schema and migration SQL support future work."],["Portability","Works on current desktop and mobile browsers."],["Integrity","Foreign keys link pledges to valid donors and campaigns."]],[1.35,5.15])
    h(d,"4. Acceptance criteria")
    numbered(d,["A donor can register, sign in, find an approved campaign and record a valid pledge.","An NGO cannot submit an active campaign until an administrator has approved its profile.","An approved, unexpired campaign appears in public results; rejected/pending campaigns do not.","An administrator can approve/reject submitted NGO profiles and campaigns.","Invalid and unauthorized mutations are rejected server-side."])
    save(d,"SRS.docx")

def testing():
    d=doc("Testing Report")
    h(d,"1. Test strategy"); p(d,"Testing combines TypeScript static checking, ESLint quality checks, production compilation, manual workflow checks, usability checks and post-deployment smoke tests. Automated tests should be expanded with Vitest and Playwright before a broader release.")
    h(d,"2. Build verification")
    table(d,["Check","Result","Evidence"],[["Type checking","Pass","npx tsc --noEmit"],["Linting","Pass","npm run lint"],["Production compilation","Pass","npx next build --webpack"],["Turbopack build","Environment limitation","Blocked by a local OS port restriction during CSS processing; webpack build passed."]],[1.7,1.0,3.8])
    h(d,"3. Functional test cases")
    rows=[["T01","Register donor","Valid donor account is created.","To execute after DB setup","Not run"],["T02","Submit NGO profile","Profile enters pending state.","To execute after DB setup","Not run"],["T03","Approve NGO","Approved NGO can submit campaign.","To execute after DB setup","Not run"],["T04","Approve campaign","Approved unexpired campaign is public.","To execute after DB setup","Not run"],["T05","Zero/negative pledge","Server validation rejects pledge.","To execute after DB setup","Not run"],["T06","Role bypass","Unauthorized server action is rejected.","To execute after DB setup","Not run"],["T07","Mobile catalogue","Search/filter controls remain usable.","To execute on deployment","Not run"]]
    table(d,["ID","Test case","Expected result","Actual result","Status"],rows,[.45,1.2,2.0,1.65,1.2])
    h(d,"4. Defects and corrective action")
    p(d,"Record any defects found during live testing here, including severity, corrective action, retest result and date. The deployment checklist requires completing T01-T07 against the final Vercel URL before LMS submission.")
    save(d,"Testing_Report.docx")

def debt():
    d=doc("Technical Debt Plan")
    h(d,"1. Purpose"); p(d,"This register records deliberate short-term compromises made to fit PROMESA within the 48-hour individual examination scope. Each item is prioritised for future repayment.")
    h(d,"2. Technical-debt register")
    rows=[["No payment gateway","Scope control","Pledges are commitments rather than transfers.","Medium","Integrate a Ghana-supported provider with signed webhooks in v2."],["No coordinator updates","Not required in MVP","Campaign progress is less detailed.","Low","Add coordinator role and milestone/update model."],["Basic reporting","Time constraint","Limited analytics and export capability.","Medium","Add date filters, exports and richer metrics."],["Manual admin seed","Small deployment model","Credential rotation is operational.","Medium","Add invite/recovery flows and audit logging."],["Limited automated workflow tests","Time constraint","Higher regression risk.","Medium","Add Vitest, Playwright and CI before expansion."],["Single-region initial deployment","MVP architecture","Potential latency/availability trade-off.","Low","Review Neon/Vercel region and monitoring needs."]]
    table(d,["Debt","Cause","Impact","Priority","Resolution"],rows,[1.2,1.15,1.55,.75,1.85])
    h(d,"3. Repayment approach"); numbered(d,["Resolve security and data-integrity debt before introducing payment collection.","Add automated regression coverage before adding coordinator or notification features.","Review the register at each release and re-prioritise against user feedback and incident data.","Document completed repayment work in release notes and architecture records."])
    save(d,"Technical_Debt_Plan.docx")

def manual():
    d=doc("User Manual")
    h(d,"1. Overview"); p(d,"PROMESA is a pledge-management application for SDG-focused donation campaigns. It records commitments in GHS; it does not process payments.")
    h(d,"2. Public visitor"); numbered(d,["Open the live PROMESA URL.","Use Campaigns to search by title or choose an SDG filter.","Open a campaign card to view its purpose, target, deadline and total pledged amount.","Register or sign in to record a pledge."])
    h(d,"3. Donor"); numbered(d,["Select Get started and choose Donor.","Enter name, email and a password of at least eight characters.","Sign in, open an approved active campaign, enter a positive GHS amount and select Record pledge.","Open the Donor dashboard to review pledge history and statuses."])
    h(d,"4. NGO representative"); numbered(d,["Register as an NGO representative and sign in.","In the NGO dashboard, submit the organisation profile and primary SDG.","Wait for the administrator to approve the profile.","Submit a campaign with title, impact description, target, deadline and SDG.","After approval, review received pledges and acknowledge or mark them fulfilled."])
    h(d,"5. Administrator"); numbered(d,["Sign in using the securely seeded administrator account.","Review pending NGO profiles and approve or reject each one.","Review pending campaign proposals and approve or reject each one.","Use the administrator dashboard to review basic platform counts and recent users."])
    h(d,"6. Troubleshooting"); table(d,["Issue","Action"],[["Cannot sign in","Confirm email/password; ask the administrator to verify that the account was created."],["Campaign is not public","Confirm NGO and campaign have both been approved and the deadline has not passed."],["Pledge cannot be recorded","Sign in as a donor and enter a positive amount for an active campaign."],["Database configuration message","Set the required Vercel/Neon environment variables and apply the migration."]],[2.0,4.5])
    save(d,"User_Manual.docx")

def links():
    d=doc("Deployment and Source Links")
    h(d,"Complete this page before submission")
    p(d,"Do not place production administrator passwords in a publicly accessible repository. Provide credentials only through the secure course submission channel if your examiner requires them.",italic=True)
    table(d,["Field","Value"],[["Student Name","[Enter your name]"],["Student ID","[Enter your student ID]"],["Project Title","PROMESA - Donation Campaign Management System for NGOs"],["Live Application","[Enter Vercel URL]"],["Admin URL","[Enter URL, normally /dashboard]"],["Test Donor Username","[Enter test account email]"],["Test Donor Password","[Enter secure test password]"],["Admin Username","[Enter admin email]"],["Admin Password","[Provide only through required secure channel]"],["Source Code Repository","[Enter GitHub repository URL]"],["Neon Database","[Do not submit connection string; record project name only if required]"],["Deployment Date","[Enter date]"]],[2.15,4.35])
    h(d,"Deployment checklist")
    numbered(d,["Apply drizzle/0000_initial.sql to the Neon database.","Add DATABASE_URL, NEXTAUTH_URL, NEXTAUTH_SECRET, ADMIN_EMAIL and ADMIN_PASSWORD in Vercel.","Run npm run seed:admin once with the secure environment variables.","Deploy the main branch through Vercel.","Run T01-T07 from the Testing Report against the live URL.","Verify all links and test credentials before uploading to Sakai LMS."])
    save(d,"Deployment_and_Source_Links.docx")

project_documentation(); srs(); testing(); debt(); manual(); links()
print(f"Created documents in {OUT}")
